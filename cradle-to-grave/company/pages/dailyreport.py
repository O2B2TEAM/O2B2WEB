import io
import os

import requests
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI

# 환경 변수 로드
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

st.set_page_config(
    page_title="노세老世",
    layout="centered",
    initial_sidebar_state="collapsed",
)

# GPT-4 모델을 호출하는 llm 객체 생성
llm = ChatOpenAI(api_key=openai_api_key)

# 프롬프트 템플릿 설정
prompt = ChatPromptTemplate.from_messages([
    ("system", "You are an AI Intern. Make a daily report."),
    ("user", "{input}")
])

# 출력 파서 생성
output_parser = StrOutputParser()

# 프롬프트 템플릿, GPT-4 모델, 출력 파서를 체인으로 연결
chain = prompt | llm | output_parser

# 뉴스 스크래핑 함수
def scrape_naver_news(query):
    url = f"https://search.naver.com/search.naver?where=news&query={query}"
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')
    articles = soup.find_all('a', {'class': 'news_tit'})
    news_list = [{'title': article.get_text(), 'link': article['href']} for article in articles]
    return news_list

st.header("노세老世 | Daily Report", divider='orange')
st.markdown("오늘의 데일리 리포트를 작성합니다.")
st.markdown(" ")

# 검색어 입력
content = st.text_input("오늘의 주제는 무엇인가요?", "")

# "작성하기" 버튼을 설정
if st.button("작성하기"):
    if not content:
        st.error("검색어를 입력하세요.")
    else:
        with st.spinner('작성 중입니다...'):
            news_list = scrape_naver_news(content)
            
            if news_list:
                news_text = '\n\n'.join([f"Title: {news['title']}\nLink: {news['link']}" for news in news_list])
                # 체인을 호출하여 요약 결과를 받아옴
                result = chain.invoke({"input": news_text + "\n\n이 내용을 요약해서 데일리 리포트를 한글로 작성해줘. 최대한 길게 작성해줘. 더 길게 작성해줘"})
                # 요약 결과를 화면에 출력
                st.markdown(f'<div style="border: 2px solid black; padding: 10px;">{result}</div>', unsafe_allow_html=True)
                
                # Word 파일 생성 함수
                def create_word_report(text):
                    doc = Document()
                    doc.add_heading('Daily Report', 0)
                    p = doc.add_paragraph(text)
                    p.style.font.size = Pt(11)
                    return doc

                # Word 파일 다운로드 버튼 생성
                doc = create_word_report(result)
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)

                st.download_button(
                    label="Word 파일 다운로드",
                    data=bio,
                    file_name="daily_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error("뉴스 기사를 찾을 수 없습니다. 다른 검색어를 시도해 보세요.")
else:
    # 버튼이 눌리지 않으면 빈 문자열 출력 (화면을 깨끗하게 유지)
    st.write("")
